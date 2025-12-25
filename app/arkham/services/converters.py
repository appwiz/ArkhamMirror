import os
import logging
import plistlib
import re
import json
import docx
import extract_msg
from email import policy
from email.parser import BytesParser
from PIL import Image
from typing import Optional, Dict, Any, List, Tuple

logger = logging.getLogger(__name__)


# =============================================================================
# TEXT-BASED FILE DETECTION
# =============================================================================
# These file types contain extractable text and don't need OCR
TEXT_BASED_EXTENSIONS = {".txt", ".eml", ".emlx", ".msg", ".docx", ".html", ".htm", ".md", ".json", ".xml", ".csv"}


def is_text_based_file(file_path: str) -> bool:
    """Check if file is text-based and doesn't need OCR."""
    ext = os.path.splitext(file_path)[1].lower()
    return ext in TEXT_BASED_EXTENSIONS


def extract_text_direct(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Extract text directly from text-based files without OCR.

    Returns:
        Dict with 'text', 'metadata' keys, or None if extraction fails.
        metadata contains any extracted headers/properties.
    """
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".txt":
            return _extract_txt(file_path)
        elif ext == ".eml":
            return _extract_eml(file_path)
        elif ext == ".emlx":
            return _extract_emlx(file_path)
        elif ext == ".msg":
            return _extract_msg_text(file_path)
        elif ext == ".docx":
            return _extract_docx_text(file_path)
        elif ext in {".html", ".htm"}:
            return _extract_html(file_path)
        elif ext == ".md":
            return _extract_md(file_path)
        elif ext == ".csv":
            return _extract_csv(file_path)
        elif ext in {".json", ".xml"}:
            return _extract_txt(file_path)  # Treat as plain text
        else:
            return None
    except Exception as e:
        logger.error(f"Direct text extraction failed for {file_path}: {e}")
        return None


def _extract_txt(file_path: str) -> Dict[str, Any]:
    """Extract text from plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return {"text": text, "metadata": {}}


def _extract_csv(file_path: str) -> Dict[str, Any]:
    """Extract CSV file as a table with markers."""
    import csv

    with open(file_path, "r", encoding="utf-8", errors="replace", newline='') as f:
        # Try to detect delimiter
        sample = f.read(4096)
        f.seek(0)

        try:
            dialect = csv.Sniffer().sniff(sample)
        except csv.Error:
            dialect = csv.excel  # Default to comma-separated

        reader = csv.reader(f, dialect)
        rows = list(reader)

    if not rows:
        return {"text": "", "metadata": {}}

    # Build table with markers
    table_lines = ["=== TABLE 1 ==="]
    for row in rows:
        table_lines.append(" | ".join(cell.strip() for cell in row))
    table_lines.append("=== END TABLE 1 ===")

    text = "\n".join(table_lines)

    return {
        "text": text,
        "metadata": {
            "row_count": len(rows),
            "col_count": len(rows[0]) if rows else 0,
        }
    }


def _extract_md(file_path: str) -> Dict[str, Any]:
    """
    Extract markdown file, converting markdown tables to our table format.

    Markdown tables look like:
    | Header1 | Header2 |
    |---------|---------|
    | Cell1   | Cell2   |
    """
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    # Find markdown tables and convert them
    # Pattern: lines starting with | and containing |
    lines = text.split('\n')
    result_lines = []
    table_buffer = []
    table_count = 0
    in_table = False

    for line in lines:
        stripped = line.strip()

        # Check if this is a table row
        is_table_row = stripped.startswith('|') and stripped.endswith('|') and '|' in stripped[1:-1]
        # Check if this is a separator row (|---|---|)
        is_separator = is_table_row and all(c in '|-: ' for c in stripped)

        if is_table_row:
            if not in_table:
                # Starting a new table
                in_table = True
                table_count += 1
                table_buffer = [f"=== TABLE {table_count} ==="]

            if not is_separator:
                # Extract cells, removing outer pipes
                cells = [cell.strip() for cell in stripped[1:-1].split('|')]
                table_buffer.append(" | ".join(cells))
        else:
            if in_table:
                # End of table
                table_buffer.append(f"=== END TABLE {table_count} ===")
                result_lines.append('\n'.join(table_buffer))
                table_buffer = []
                in_table = False

            result_lines.append(line)

    # Handle table at end of file
    if in_table:
        table_buffer.append(f"=== END TABLE {table_count} ===")
        result_lines.append('\n'.join(table_buffer))

    return {"text": '\n'.join(result_lines), "metadata": {}}


def _extract_eml(file_path: str) -> Dict[str, Any]:
    """Extract text and metadata from .eml file."""
    with open(file_path, "rb") as f:
        msg = BytesParser(policy=policy.default).parse(f)

    # Extract body
    body_part = msg.get_body(preferencelist=("plain", "html"))
    if body_part:
        body = body_part.get_content()
        # Strip HTML if needed
        if body_part.get_content_type() == "text/html":
            body = _strip_html(body)
    else:
        body = ""

    # Extract metadata
    metadata = {
        "email_subject": str(msg.get("subject", "")),
        "email_from": str(msg.get("from", "")),
        "email_to": str(msg.get("to", "")),
        "email_cc": str(msg.get("cc", "")),
        "email_date": str(msg.get("date", "")),
        "email_message_id": str(msg.get("message-id", "")),
        "email_in_reply_to": str(msg.get("in-reply-to", "")),
    }

    # Build full text with headers for search
    header_text = f"""Subject: {metadata['email_subject']}
From: {metadata['email_from']}
To: {metadata['email_to']}
CC: {metadata['email_cc']}
Date: {metadata['email_date']}

"""

    return {"text": header_text + body, "metadata": metadata}


def _extract_emlx(file_path: str) -> Dict[str, Any]:
    """
    Extract text and metadata from .emlx file (Apple Mail format).

    EMLX format:
    - First line: byte count of the message
    - Followed by: RFC 822 email message
    - Followed by: XML plist with Apple Mail metadata
    """
    with open(file_path, "rb") as f:
        content = f.read()

    # First line is the byte count
    first_newline = content.find(b"\n")
    if first_newline == -1:
        raise ValueError("Invalid emlx format: no newline found")

    try:
        byte_count = int(content[:first_newline].decode("ascii").strip())
    except ValueError:
        # Some emlx files don't have the byte count, treat as raw email
        byte_count = len(content)
        first_newline = -1

    # Extract the email portion
    email_start = first_newline + 1
    email_end = email_start + byte_count
    email_data = content[email_start:email_end]

    # Parse the email
    msg = BytesParser(policy=policy.default).parsebytes(email_data)

    # Extract body
    body_part = msg.get_body(preferencelist=("plain", "html"))
    if body_part:
        body = body_part.get_content()
        if body_part.get_content_type() == "text/html":
            body = _strip_html(body)
    else:
        body = ""

    # Extract email metadata
    metadata = {
        "email_subject": str(msg.get("subject", "")),
        "email_from": str(msg.get("from", "")),
        "email_to": str(msg.get("to", "")),
        "email_cc": str(msg.get("cc", "")),
        "email_date": str(msg.get("date", "")),
        "email_message_id": str(msg.get("message-id", "")),
        "email_in_reply_to": str(msg.get("in-reply-to", "")),
    }

    # Try to extract Apple Mail plist metadata
    try:
        plist_start = content.find(b"<?xml", email_end)
        if plist_start != -1:
            plist_data = content[plist_start:]
            apple_meta = plistlib.loads(plist_data)
            # Common Apple Mail metadata fields
            if "flags" in apple_meta:
                metadata["apple_flags"] = apple_meta["flags"]
            if "date-received" in apple_meta:
                metadata["email_received_date"] = str(apple_meta["date-received"])
    except Exception as e:
        logger.debug(f"Could not parse Apple Mail plist: {e}")

    # Build full text with headers
    header_text = f"""Subject: {metadata['email_subject']}
From: {metadata['email_from']}
To: {metadata['email_to']}
CC: {metadata['email_cc']}
Date: {metadata['email_date']}

"""

    return {"text": header_text + body, "metadata": metadata}


def _extract_msg_text(file_path: str) -> Dict[str, Any]:
    """Extract text and metadata from .msg file (Outlook format)."""
    msg = extract_msg.Message(file_path)

    metadata = {
        "email_subject": msg.subject or "",
        "email_from": msg.sender or "",
        "email_to": msg.to or "",
        "email_cc": msg.cc or "",
        "email_date": str(msg.date) if msg.date else "",
    }

    header_text = f"""Subject: {metadata['email_subject']}
From: {metadata['email_from']}
To: {metadata['email_to']}
CC: {metadata['email_cc']}
Date: {metadata['email_date']}

"""

    body = msg.body or ""
    msg.close()

    return {"text": header_text + body, "metadata": metadata}


def _extract_docx_text(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Extract text from .docx file including tables.

    Tables are wrapped with markers for chunking protection:
    === TABLE N ===
    Header1 | Header2 | Header3
    Row1Col1 | Row1Col2 | Row1Col3
    === END TABLE N ===
    """
    logger.info(f"Starting DOCX extraction: {file_path}")

    # Validate file exists
    if not os.path.exists(file_path):
        logger.error(f"DOCX file not found: {file_path}")
        return None

    try:
        file_size = os.path.getsize(file_path)
        logger.info(f"DOCX file size: {file_size:,} bytes")
    except Exception as stat_err:
        logger.warning(f"Could not get file stats: {stat_err}")

    try:
        doc = docx.Document(file_path)
        logger.info(f"DOCX opened successfully. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

        text_parts = []

        # Extract all paragraph text
        para_count = 0
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                para_count += 1
        logger.info(f"Extracted {para_count} non-empty paragraphs")

        # Extract tables
        table_count = 0
        for table_idx, table in enumerate(doc.tables):
            try:
                table_lines = []
                table_lines.append(f"\n=== TABLE {table_idx + 1} ===")

                row_count = 0
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    # Deduplicate merged cells (Word tables can repeat content for merged cells)
                    unique_cells = []
                    prev_cell = None
                    for cell in cells:
                        if cell != prev_cell:
                            unique_cells.append(cell)
                            prev_cell = cell
                    table_lines.append(" | ".join(unique_cells))
                    row_count += 1

                table_lines.append(f"=== END TABLE {table_idx + 1} ===\n")
                text_parts.append("\n".join(table_lines))
                table_count += 1
                logger.debug(f"Extracted table {table_idx + 1} with {row_count} rows")

            except Exception as table_err:
                logger.warning(f"Failed to extract table {table_idx + 1}: {table_err}")
                continue

        logger.info(f"Extracted {table_count} tables successfully")

        text = "\n".join(text_parts)
        logger.info(f"Total extracted text length: {len(text)} characters")

        # Extract core properties as metadata
        metadata = {}
        try:
            props = doc.core_properties
            if props.author:
                metadata["doc_author"] = props.author
            if props.title:
                metadata["doc_title"] = props.title
            if props.subject:
                metadata["doc_subject"] = props.subject
            if props.created:
                metadata["doc_created"] = str(props.created)
            if props.modified:
                metadata["doc_modified"] = str(props.modified)
        except Exception as e:
            logger.debug(f"Could not extract docx properties: {e}")

        return {"text": text, "metadata": metadata}

    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {e}", exc_info=True)
        return None


def _extract_html(file_path: str) -> Dict[str, Any]:
    """Extract text from HTML file, including tables with markers."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        html = f.read()

    text_parts = []

    # Extract tables first, before stripping HTML
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, 'lxml')

        # Extract tables with markers
        for table_idx, table in enumerate(soup.find_all('table')):
            table_lines = [f"\n=== TABLE {table_idx + 1} ==="]

            rows = table.find_all('tr')
            for row in rows:
                cells = row.find_all(['th', 'td'])
                cell_texts = [cell.get_text(strip=True) for cell in cells]
                if any(cell_texts):  # Skip empty rows
                    table_lines.append(" | ".join(cell_texts))

            table_lines.append(f"=== END TABLE {table_idx + 1} ===\n")

            if len(table_lines) > 2:  # Has actual content
                text_parts.append("\n".join(table_lines))

            # Remove table from soup to avoid duplicate extraction
            table.decompose()

        # Get remaining text (non-table content)
        remaining_text = soup.get_text(separator=' ', strip=True)
        if remaining_text:
            text_parts.insert(0, remaining_text)  # Put main text first

    except ImportError:
        # Fallback if BeautifulSoup not available (it should be - it's in requirements)
        logger.warning("BeautifulSoup not available, falling back to simple HTML strip")
        text_parts.append(_strip_html(html))

    return {"text": "\n\n".join(text_parts), "metadata": {}}


def _strip_html(html: str) -> str:
    """Simple HTML tag stripper."""
    # Remove script and style elements
    html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL | re.IGNORECASE)
    # Remove all tags
    html = re.sub(r'<[^>]+>', ' ', html)
    # Clean up whitespace
    html = re.sub(r'\s+', ' ', html)
    return html.strip()


# =============================================================================
# TABLE EXTRACTION FROM TEXT
# =============================================================================

def extract_tables_from_text(text: str) -> List[Dict[str, Any]]:
    """
    Parse table markers from extracted text and return table metadata.

    Looks for patterns like:
    === TABLE N ===
    Header1 | Header2 | Header3
    Row1Col1 | Row1Col2 | Row1Col3
    === END TABLE N ===

    Returns:
        List of dicts with keys: table_index, row_count, col_count, headers, text_content
    """
    tables = []

    # Pattern to match table blocks
    table_pattern = re.compile(
        r'=== TABLE (\d+) ===\s*\n(.*?)=== END TABLE \1 ===',
        re.DOTALL
    )

    for match in table_pattern.finditer(text):
        table_index = int(match.group(1))
        table_content = match.group(2).strip()

        if not table_content:
            continue

        lines = [line.strip() for line in table_content.split('\n') if line.strip()]
        if not lines:
            continue

        # First line is headers
        headers = [cell.strip() for cell in lines[0].split('|')]

        # Count rows and columns
        row_count = len(lines)
        col_count = len(headers)

        # Full text representation
        text_content = f"=== TABLE {table_index} ===\n{table_content}\n=== END TABLE {table_index} ==="

        tables.append({
            "table_index": table_index,
            "page_num": 1,  # Text files are treated as single page
            "row_count": row_count,
            "col_count": col_count,
            "headers": json.dumps(headers),
            "text_content": text_content,
        })

    return tables


# =============================================================================
# PDF CONVERSION (for files that need visual representation)
# =============================================================================

def convert_to_pdf(file_path):
    """
    Converts various file formats (.docx, .msg, .eml, .emlx, images) to PDF.
    Returns the path to the generated PDF.

    NOTE: For text-based files, consider using extract_text_direct() instead
    to skip the PDF->OCR round-trip.
    """
    ext = os.path.splitext(file_path)[1].lower()
    output_pdf_path = file_path + ".converted.pdf"

    try:
        if ext == ".docx":
            _convert_docx_to_pdf(file_path, output_pdf_path)
        elif ext == ".msg":
            _convert_msg_to_pdf(file_path, output_pdf_path)
        elif ext == ".eml":
            _convert_eml_to_pdf(file_path, output_pdf_path)
        elif ext == ".emlx":
            _convert_emlx_to_pdf(file_path, output_pdf_path)
        elif ext == ".txt":
            _convert_txt_to_pdf(file_path, output_pdf_path)
        elif ext in {".html", ".htm"}:
            _convert_html_to_pdf(file_path, output_pdf_path)
        elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
            _convert_image_to_pdf(file_path, output_pdf_path)
        else:
            raise ValueError(f"Unsupported file type for conversion: {ext}")

        return output_pdf_path
    except Exception as e:
        logger.error(f"Conversion failed for {file_path}: {e}")
        if os.path.exists(output_pdf_path):
            os.remove(output_pdf_path)
        raise e


def _convert_docx_to_pdf(docx_path, pdf_path):
    # Quick and dirty: Extract text and create a simple PDF using ReportLab
    # Ideally, we'd use LibreOffice or win32com for perfect layout preservation,
    # but that introduces heavy dependencies. For v0.1, text extraction is key.
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    doc = docx.Document(docx_path)
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y = height - 40

    for para in doc.paragraphs:
        text = para.text
        # Simple wrapping
        lines = _simple_wrap(text, 80)
        for line in lines:
            if y < 40:
                c.showPage()
                y = height - 40
            c.drawString(40, y, line)
            y -= 12

    c.save()


def _convert_msg_to_pdf(msg_path, pdf_path):
    msg = extract_msg.Message(msg_path)
    _create_text_pdf(
        msg.body,
        pdf_path,
        f"Subject: {msg.subject}\nFrom: {msg.sender}\nTo: {msg.to}\nDate: {msg.date}\n\n",
    )
    msg.close()


def _convert_eml_to_pdf(eml_path, pdf_path):
    with open(eml_path, "rb") as f:
        msg = BytesParser(policy=policy.default).parse(f)

    body_part = msg.get_body(preferencelist=("plain", "html"))
    if body_part:
        body = body_part.get_content()
        if body_part.get_content_type() == "text/html":
            body = _strip_html(body)
    else:
        body = ""
    header = f"Subject: {msg['subject']}\nFrom: {msg['from']}\nTo: {msg['to']}\nDate: {msg['date']}\n\n"
    _create_text_pdf(body, pdf_path, header)


def _convert_emlx_to_pdf(emlx_path, pdf_path):
    """Convert .emlx (Apple Mail) file to PDF."""
    extracted = _extract_emlx(emlx_path)
    _create_text_pdf(extracted["text"], pdf_path)


def _convert_html_to_pdf(html_path, pdf_path):
    """Convert HTML file to PDF."""
    extracted = _extract_html(html_path)
    _create_text_pdf(extracted["text"], pdf_path)


def _convert_image_to_pdf(img_path, pdf_path):
    image = Image.open(img_path)
    if image.mode != "RGB":
        image = image.convert("RGB")
    image.save(pdf_path, "PDF", resolution=100.0)


def _convert_txt_to_pdf(txt_path, pdf_path):
    with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    _create_text_pdf(text, pdf_path)


def _create_text_pdf(text, pdf_path, header=""):
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y = height - 40

    full_text = header + (text or "")

    for line in full_text.split("\n"):
        wrapped_lines = _simple_wrap(line, 90)
        for w_line in wrapped_lines:
            if y < 40:
                c.showPage()
                y = height - 40
            c.drawString(40, y, w_line)
            y -= 12
    c.save()


def _simple_wrap(text, max_chars):
    return [text[i : i + max_chars] for i in range(0, len(text), max_chars)]
